import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
import time
import fitz  # PyMuPDF for reading PDFs
import docx  # For DOCX files
import re
import io
import base64
import pandas as pd
import nltk
from nltk.tokenize import sent_tokenize
import hashlib
from datetime import datetime
import plotly.express as px
import plotly.graph_objects as go
import toml
import os

# Create directory for NLTK data if it doesn't exist
nltk_data_dir = os.path.join(os.getcwd(), 'nltk_data')
os.makedirs(nltk_data_dir, exist_ok=True)

# Set custom path to NLTK data directory
nltk.data.path.append(nltk_data_dir)

# Download NLTK data (first time only)
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', download_dir=nltk_data_dir)

# Page configuration with more options
st.set_page_config(
    page_title="Advanced AI-Powered Text Summarizer",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize theme in session state if not exists
if 'theme' not in st.session_state:
    st.session_state.theme = "Light"

# Apply theme based on session state
if st.session_state.theme == "Dark":
    # Dark theme
    st.markdown("""
    <style>
    body {background-color: #121212; color: white;}
    .main {background-color: #121212; color: white;}
    .stTextInput > div > div > input {background-color: #2E2E2E; color: white; border-color: #555555;}
    .stTextArea > div > div > textarea {background-color: #2E2E2E; color: white; border-color: #555555;}
    .stSelectbox > div > div > select {background-color: #2E2E2E; color: white;}
    .stSelectbox > div > div {background-color: #2E2E2E; color: white;}
    .stButton > button {border-color: #4CAF50; color: white; background-color: #4CAF50;}
    .reportview-container {background-color: #121212;}
    .sidebar .sidebar-content {background-color: #1E1E1E; color: white;}
    h1, h2, h3, h4, h5, h6 {color: white !important;}
    .stMarkdown {color: white;}
    .stDataFrame {color: white;}
    .css-145kmo2 {color: white !important;}
    .css-1aumxhk {background-color: #1E1E1E !important;}
    </style>
    """, unsafe_allow_html=True)
else:
    # Light theme (default)
    st.markdown("""
    <style>
    .main {background-color: white; color: #262730;}
    </style>
    """, unsafe_allow_html=True)

# Custom CSS for better styling (applied regardless of theme)
# Apply theme based on session state
if st.session_state.theme == "Dark":
    # Dark theme with specific sidebar targeting
    st.markdown("""
    <style>
    /* Main containers */
    .main {background-color: #121212 !important; color: white !important;}
    .stApp {background-color: #121212 !important;}
    .css-18e3th9 {background-color: #121212 !important; color: white !important;}
    .css-1d391kg {background-color: #121212 !important;}

    /* Sidebar - target all possible sidebar selectors */
    .css-1e5imcs {background-color: #1E1E1E !important;}
    .css-hxt7ib {background-color: #1E1E1E !important;}
    section[data-testid="stSidebar"] {background-color: #1E1E1E !important;}
    .st-emotion-cache-16txtl3 {background-color: #1E1E1E !important;}
    div[data-testid="stSidebarNav"] {background-color: #1E1E1E !important;}
    .css-16idsys p {color: white !important;}
    .css-16idsys {color: white !important;}
    .st-emotion-cache-16idsys p {color: white !important;}
    .st-emotion-cache-16idsys {color: white !important;}

    /* Text inputs */
    .stTextInput > div > div > input {background-color: #2E2E2E !important; color: white !important; border-color: #555555 !important;}
    .stTextArea > div > div > textarea {background-color: #2E2E2E !important; color: white !important; border-color: #555555 !important;}
    .stSelectbox > div > div > select {background-color: #2E2E2E !important; color: white !important;}
    .stSelectbox > div > div {background-color: #2E2E2E !important; color: white !important;}

    /* Buttons */
    .stButton > button {border-color: #4CAF50 !important; color: white !important; background-color: #4CAF50 !important;}

    /* Text elements */
    h1, h2, h3, h4, h5, h6 {color: white !important;}
    .stMarkdown {color: white !important;}
    p {color: white !important;}
    label {color: white !important;}

    /* Data elements */
    .stDataFrame {color: white !important;}
    .dataframe {color: white !important;}

    /* Other elements */
    .css-145kmo2 {color: white !important;}
    .css-1aumxhk {background-color: #1E1E1E !important;}
    .reportview-container {background-color: #121212 !important;}

    /* Additional sidebar specific selectors to ensure they're targeted */
    [data-testid="stSidebar"] [data-testid="stMarkdown"] p {color: white !important;}
    [data-testid="stSidebar"] [data-testid="stMarkdown"] {color: white !important;}
    [data-testid="stSidebar"] div {color: white !important;}
    [data-testid="stSidebar"] h1 {color: white !important;}
    [data-testid="stSidebar"] h2 {color: white !important;}
    [data-testid="stSidebar"] a {color: #4CAF50 !important;}
    </style>
    """, unsafe_allow_html=True)
else:
    # Light theme (default)
    st.markdown("""
    <style>
    /* Light theme resets */
    .main {background-color: white !important; color: #262730 !important;}
    .stApp {background-color: white !important;}
    .footer {
        margin-top: 3rem;
        text-align: center;
        color: #757575;
        font-size: 0.8rem;
    }
    </style>
    """, unsafe_allow_html=True)

load_dotenv()

# Load API Key 
try:
    # Method 1: Try environment variables first
    api_key = os.environ.get("OPEN_API_KEY")

    # Method 2: Try Streamlit secrets second
    if not api_key:
        try:
            api_key = st.secrets["OPEN_API_KEY"]
        except:
            pass

    # Initialize client if we found an API key
    if api_key:
        client = OpenAI(api_key=api_key)
        api_loaded = True
    else:
        api_loaded = False
        st.error("API key not found. Please configure OPEN_API_KEY in environment variables or Streamlit secrets.")

except Exception as e:
    st.error(f"Error loading API key: {e}")
    st.info("Please set OPEN_API_KEY in your deployment environment")
    api_loaded = False

# Initialize session state
if 'summaries' not in st.session_state:
    st.session_state.summaries = []
if 'current_text' not in st.session_state:
    st.session_state.current_text = ""
if 'current_summary' not in st.session_state:
    st.session_state.current_summary = ""
if 'summary_count' not in st.session_state:
    st.session_state.summary_count = 0
if 'processing_time' not in st.session_state:
    st.session_state.processing_time = []
if 'char_counts' not in st.session_state:
    st.session_state.char_counts = []

# Sidebar for settings and history
with st.sidebar:
    try:
        st.image("Image/RCB LOGO.png", width=200)
    except:
        st.title("AI-Powered Text Summarizer")

    st.title("Settings")

    # Theme toggle
    theme = st.radio("Theme", ["Light", "Dark"])
    if theme != st.session_state.theme:
        st.session_state.theme = theme
        st.rerun()

    # Model selection
    model = st.selectbox(
        "Select AI Model",
        ["gpt-3.5-turbo", "gpt-4"],
        index=1,
        help="GPT-4 provides higher quality summaries but is slightly slower"
    )

    # Advanced settings
    with st.expander("Advanced Settings"):
        temperature = st.slider(
            "Temperature",
            min_value=0.0,
            max_value=1.0,
            value=0.5,
            step=0.1,
            help="Higher values make output more creative, lower values more deterministic"
        )

        max_tokens = st.slider(
            "Max Output Length",
            min_value=50,
            max_value=3000,
            value=600,
            step=50,
            help="Maximum length of the summary"
        )

        animate_text = st.checkbox(
            "Animate Summary Text",
            value=True,
            help="Show typing animation when displaying the summary"
        )

    # History section
    st.header("Summary History")
    if st.session_state.summaries:
        for i, summary_data in enumerate(st.session_state.summaries):
            with st.expander(f"Summary {i + 1}: {summary_data['date']}"):
                st.write(f"**Type:** {summary_data['type']}")
                st.write(f"**Model:** {summary_data['model']}")
                st.write(f"**Text Sample:** {summary_data['text'][:100]}...")
                st.write("**Summary:**")
                st.write(summary_data['summary'])
                if st.button("Use This Text Again", key=f"use_again_{i}"):
                    st.session_state.current_text = summary_data['full_text']
                    st.rerun()
    else:
        st.info("Your summary history will appear here")

        # Analytics section
        if st.session_state.summary_count > 0:
            st.header("Analytics")

            # Processing time chart
            if st.session_state.processing_time:
                fig = px.line(
                    x=list(range(1, len(st.session_state.processing_time) + 1)),
                    y=st.session_state.processing_time,
                    title="Processing Time",
                    labels={"x": "Summary #", "y": "Time (seconds)"}
                )
                st.plotly_chart(fig, use_container_width=True)

            # Character count chart
            if st.session_state.char_counts:
                fig = px.bar(
                    x=list(range(1, len(st.session_state.char_counts) + 1)),
                    y=st.session_state.char_counts,
                    title="Document Size",
                    labels={"x": "Summary #", "y": "Characters"}
                )
                st.plotly_chart(fig, use_container_width=True)

# Main Content Area
st.markdown('<div class="header-container">', unsafe_allow_html=True)
st.markdown('<h1 class="main-header">AI-Powered Text Summarizer</h1>', unsafe_allow_html=True)
st.markdown('</div>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">Upload a document or paste text to generate an AI-powered summary</p>',
            unsafe_allow_html=True)

# Create tabs for different functionalities
tab1, tab2, tab3 = st.tabs(["Summarize", "Compare Documents", "Help & Examples"])

# Tab 1: Main summarization functionality
with tab1:
    # Summary type selection with descriptions
    st.subheader("Summary Type")
    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("""
        <div style="border: 1px solid #1E88E5; border-radius: 0.5rem; padding: 1rem; height: 150px;">
            <h3 style="color: #1E88E5;">Short</h3>
            <p>2-3 bullet points highlighting only the most critical information</p>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        st.markdown("""
        <div style="border: 1px solid #1E88E5; border-radius: 0.5rem; padding: 1rem; height: 150px;">
            <h3 style="color: #1E88E5;">Medium</h3>
            <p>3-4 sentence executive summary covering main points</p>
        </div>
        """, unsafe_allow_html=True)

    with col3:
        st.markdown("""
        <div style="border: 1px solid #1E88E5; border-radius: 0.5rem; padding: 1rem; height: 150px;">
            <h3 style="color: #1E88E5;">Detailed</h3>
            <p>Comprehensive summary with all key points and insights</p>
        </div>
        """, unsafe_allow_html=True)

    summary_type = st.select_slider(
        "Select summary type",
        options=["Short", "Medium", "Detailed"],
        value="Medium"
    )

    # Custom prompt option
    custom_prompt_enabled = st.checkbox("Use custom prompt template")
    custom_prompt = ""

    if custom_prompt_enabled:
        template_options = {
            "Default": "Summarize the following text.",
            "Data Science": "Summarize the following text, focusing on key data points, methodologies, insights, and statistical trends relevant to data analysis.",
            "Health": "Provide a healthcare summary, focusing on patient data, medical insights, and any relevant treatments or interventions discussed in the text.",
            "Legal": "Provide a legal analysis of the following text, highlighting key clauses, terms, legal precedents, and potential legal implications.",
            "Finance": "Summarize the following text with a focus on financial aspects, including market trends, revenue, investment strategies, and economic implications.",
            "Custom": "Write your own prompt to tailor the summary to your specific needs."
        }

        template_choice = st.selectbox("Select a prompt template", list(template_options.keys()))

        if template_choice == "Custom":
            custom_prompt = st.text_area(
                "Enter your custom prompt template:",
                value="Summarize the following text focusing on...",
                height=100
            )
        else:
            custom_prompt = template_options[template_choice]

    # Input method selection
    st.subheader("Input Method")
    input_option = st.radio("Choose input method:", ["Paste text", "Upload Document"])

    text = ""

    # Input: Paste text
    if input_option == "Paste text":
        text = st.text_area(
            "Enter your text here:",
            height=250,
            value=st.session_state.current_text
        )
        st.session_state.current_text = text

        # Show character count for text input
        if text:
            word_count = len(re.findall(r'\w+', text))
            st.info(f"Character count: {len(text)} | Word count: {word_count} | Estimated tokens: {len(text) // 4}")

    # Input: Upload document
    elif input_option == "Upload Document":
        uploaded_file = st.file_uploader(
            "Upload a file",
            type=["pdf", "txt", "docx", "csv"],
            help="Supported formats: PDF, TXT, DOCX, CSV"
        )

        if uploaded_file is not None:
            # Display file details
            file_details = {
                "Filename": uploaded_file.name,
                "File type": uploaded_file.type,
                "File size": f"{uploaded_file.size / 1024:.2f} KB"
            }

            st.write("File Details:")
            st.json(file_details)

            # Process different file types
            with st.spinner("Processing document..."):
                try:
                    if uploaded_file.type == "application/pdf":
                        with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
                            total_pages = len(doc)
                            st.info(f"PDF document with {total_pages} pages")

                            # Progress bar for PDF processing
                            progress_bar = st.progress(0)

                            for i, page in enumerate(doc):
                                text += page.get_text()
                                progress = (i + 1) / total_pages
                                progress_bar.progress(progress)

                            # Extract tables if present
                            tables_text = ""
                            for page in doc:
                                tables = page.find_tables()
                                if tables.tables:
                                    tables_text += "\n\nTables found in document:\n"
                                    for i, table in enumerate(tables.tables):
                                        tables_text += f"\nTable {i + 1}:\n"
                                        tables_text += str(table.extract())

                            if tables_text:
                                text += tables_text

                    elif uploaded_file.type == "text/plain":
                        text = uploaded_file.getvalue().decode("utf-8")

                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
                        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

                    elif uploaded_file.type == "text/csv":
                        df = pd.read_csv(uploaded_file)
                        text = f"CSV Data Summary:\n\n"
                        text += f"Rows: {df.shape[0]}, Columns: {df.shape[1]}\n\n"
                        text += f"Column Names: {', '.join(df.columns.tolist())}\n\n"
                        text += f"Data Sample:\n{df.head(5).to_string()}\n\n"
                        text += f"Statistical Summary:\n{df.describe().to_string()}"

                    else:
                        st.error("This file type is not supported yet.")
                        text = ""

                    st.session_state.current_text = text

                    # Show text preview and stats
                    if text:
                        word_count = len(re.findall(r'\w+', text))

                        st.info(f"Successfully extracted {len(text)} characters, {word_count} words")

                        with st.expander("Show text preview"):
                            st.text_area("Document Content Preview", text[:2000] + ("..." if len(text) > 2000 else ""),
                                         height=300)

                except Exception as e:
                    st.error(f"Error processing document: {str(e)}")
                    text = ""


    # Function to chunk text for large documents
    def chunk_text(text, max_chars=4000):
        """Split text into chunks of approximately max_chars, respecting sentence boundaries"""
        if len(text) <= max_chars:
            return [text]

        sentences = sent_tokenize(text)
        chunks = []
        current_chunk = ""

        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= max_chars:
                current_chunk += " " + sentence
            else:
                chunks.append(current_chunk.strip())
                current_chunk = sentence

        if current_chunk:
            chunks.append(current_chunk.strip())

        return chunks


    # Function to summarize using OpenAI
    def get_summary(text, summary_type, model= None, temp= None, max_output_tokens= None, custom_prompt=None):
        model = model or st.session_state.model  # Use passed model if provided, else use session_state value
        temp = temp or st.session_state.temperature  # Use passed temperature if provided, else use session_state value
        max_output_tokens = max_output_tokens or st.session_state.max_tokens
        """Generate summary using OpenAI API"""
        if not api_loaded:
            return "API key not loaded. Please check your configuration.", 0

        start_time = time.time()

        prompt_styles = {
            "Short": "Summarize the following text in 1-3 bullet points covering only the most critical information.",
            "Medium": "Write a 3-4 sentence executive summary of the following text covering the main points.",
            "Detailed": "Write a detailed summary covering all key insights, information, and action points from the following text."
        }

        # Use custom prompt if provided
        prompt_template = custom_prompt if custom_prompt else prompt_styles[summary_type]

        # Check if text needs to be chunked
        chunks = chunk_text(text)

        if len(chunks) == 1:
            # Single chunk processing
            prompt = f"{prompt_template}\n\n{text}"

            try:
                response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system",
                         "content": "You are a professional summarizer that creates clear, concise, and accurate summaries."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=temp,
                    max_tokens=max_output_tokens
                )

                summary = response.choices[0].message.content.strip()

            except Exception as e:
                st.error(f"API Error: {str(e)}")
                return f"Error generating summary: {str(e)}", 0

        else:
            # Multi-chunk processing for large documents
            st.info(f"Document is large, processing in {len(chunks)} chunks...")
            progress_bar = st.progress(0)

            chunk_summaries = []

            for i, chunk in enumerate(chunks):
                chunk_prompt = f"Summarize the following section of a larger document:\n\n{chunk}"

                try:
                    response = client.chat.completions.create(
                        model=model,
                        messages=[
                            {"role": "system",
                             "content": "You are a professional summarizer that creates clear, concise, and accurate summaries."},
                            {"role": "user", "content": chunk_prompt}
                        ],
                        temperature=temp,
                        max_tokens=max_output_tokens // 2
                    )

                    chunk_summary = response.choices[0].message.content.strip()
                    chunk_summaries.append(chunk_summary)

                    # Update progress
                    progress_bar.progress((i + 1) / len(chunks))

                except Exception as e:
                    st.error(f"API Error on chunk {i + 1}: {str(e)}")
                    chunk_summaries.append(f"[Error summarizing chunk {i + 1}]")

            # Combine and summarize the chunk summaries
            combined_text = "\n\n".join(chunk_summaries)
            final_prompt = f"{prompt_template}\n\nBelow are summaries of different sections of a document. Create a coherent overall summary:\n\n{combined_text}"

            try:
                response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system",
                         "content": "You are a professional summarizer that creates clear, concise, and accurate summaries."},
                        {"role": "user", "content": final_prompt}
                    ],
                    temperature=temp,
                    max_tokens=max_output_tokens
                )

                summary = response.choices[0].message.content.strip()

            except Exception as e:
                st.error(f"API Error in final summarization: {str(e)}")
                return combined_text, 0

        processing_time = time.time() - start_time

        return summary, processing_time


    # Create download button for text
    def get_download_link(text, filename, link_text):
        """Generate a download link for text content"""
        b64 = base64.b64encode(text.encode()).decode()
        href = f'<a href="data:text/plain;base64,{b64}" download="{filename}">{link_text}</a>'
        return href

    # Validate custom prompt
    def validate_custom_prompt_with_model(prompt_type, text, model="gpt-3.5-turbo"):
        """
        Use an AI model to validate if the text matches the selected prompt domain
        Returns: Tuple of (is_valid, error_message)
        """
        if prompt_type == "Default" or prompt_type == "Custom":
            # Default and custom prompts are always valid
            return True, ""

        # Create a shortened version of the text for validation (to save on API costs)
        # Get first 300 chars and a sample from the middle
        text_sample = text[:300]
        if len(text) > 600:
            text_sample += "... " + text[len(text) // 2 - 150:len(text) // 2 + 150] + "..."

        # Make prompt to check domain match
        validation_prompt = f"""
        I need to determine if the following text is appropriate for a {prompt_type} analysis/summary.

        Text: 
        {text_sample}

        Is this text related to {prompt_type}? Respond with ONLY "YES" if it contains sufficient relevant {prompt_type} content, or "NO" if it doesn't match the {prompt_type} domain.
        """

        try:
            # Add a timeout to prevent hanging
            import time
            start_time = time.time()
            max_time = 5  # 5 seconds timeout

            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are a content validator. Respond with ONLY 'YES' or 'NO'."},
                    {"role": "user", "content": validation_prompt}
                ],
                temperature=0.1,  # Low temperature for more deterministic results
                max_tokens=5  # We only need a short response
            )

            elapsed_time = time.time() - start_time
            if elapsed_time > max_time:
                # Timed out, allow it to proceed
                return True, f"Note: Validation took too long ({elapsed_time:.1f}s), proceeding anyway."

            result = response.choices[0].message.content.strip().upper()

            if "YES" in result:
                return True, ""
            else:
                domain_descriptions = {
                    "Data Science": "data science, analytics, or statistics",
                    "Health": "healthcare, medical information, or patient care",
                    "Legal": "legal documents, contracts, or legal analysis",
                    "Finance": "financial data, market analysis, or economic information"
                }
                domain_desc = domain_descriptions.get(prompt_type, prompt_type.lower())
                return False, f"The selected text doesn't appear to contain sufficient {domain_desc} content for a {prompt_type} summary. Please select a different prompt template that fits the domain or use the Default template for general summary."

        except Exception as e:
            # If there's an API error, default to allowing it (don't block the user)
            return True, f"Note: Couldn't validate prompt match ({str(e)}), proceeding anyway."


    # Summarize button
    if st.button("Generate Summary", type="primary"):
        if not api_loaded:
            st.error("API key not loaded. Please check your configuration in KEYS/secrets.toml")
        elif text.strip() == "":
            st.warning("Please enter or upload text to summarize.")
        else:
            st.session_state.current_text = text  # Save current text

            # Validate custom prompt if enabled
            if custom_prompt_enabled and template_choice != "Default" and template_choice != "Custom":
                with st.spinner("Validating content compatibility..."):
                    is_valid, error_message = validate_custom_prompt_with_model(template_choice, text)
                    if not is_valid:
                        st.error(error_message)
                        st.stop()

            with st.spinner("Generating your summary..."):
                # Get the summary with processing time
                final_prompt = custom_prompt if custom_prompt_enabled else None
                summary, proc_time = get_summary(
                    text,
                    summary_type,
                    model=model,
                    temp=temperature,
                    max_output_tokens=max_tokens,
                    custom_prompt=final_prompt
                )


            if "Error" in summary or proc_time == 0:
                st.error(summary)
            else:
                # Display success message
                st.success(f"Summary generated in {proc_time:.2f} seconds!")

                # Save to session state for analytics
                st.session_state.processing_time.append(proc_time)
                st.session_state.char_counts.append(len(text))

                # Save summary to history
                summary_id = hashlib.md5(text.encode()).hexdigest()[:8]
                st.session_state.summaries.append({
                    'id': summary_id,
                    'date': datetime.now().strftime("%Y-%m-%d %H:%M"),
                    'type': summary_type,
                    'model': model,
                    'text': text[:100],
                    'full_text': text,
                    'summary': summary,
                    'processing_time': proc_time
                })

                st.session_state.summary_count += 1
                st.session_state.current_summary = summary

                # Create a container for the summary with styling
                summary_container_class = "summary-container"
                if st.session_state.theme == "Dark":
                    summary_container_class += " dark-summary-container"

                st.markdown(f'<div class="{summary_container_class}">', unsafe_allow_html=True)

                # Display metrics about the summary
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Original Length", f"{len(text):,} chars")
                with col2:
                    st.metric("Summary Length", f"{len(summary):,} chars")
                with col3:
                    compression = round(100 - (len(summary) / len(text) * 100), 1)
                    st.metric("Compression", f"{compression}%")

                # Display the summary with or without animation
                placeholder = st.empty()

                if animate_text and len(summary) < 2000:  # Only animate shorter summaries
                    animated_text = ""
                    # Faster animation for longer text
                    animation_speed = max(0.001, min(0.01, 2 / len(summary)))

                    for char in summary:
                        animated_text += char
                        placeholder.text_area("Summary:", animated_text, height=300)
                        time.sleep(animation_speed)
                else:
                    placeholder.text_area("Summary:", summary, height=300)

                # Create download links
                st.markdown("### Download Options")
                col1, col2 = st.columns(2)

                with col1:
                    st.markdown(
                        get_download_link(summary, f"summary_{summary_id}.txt", "Download Summary"),
                        unsafe_allow_html=True
                    )

                with col2:
                    st.markdown(
                        get_download_link(text, f"original_{summary_id}.txt", "Download Original Text"),
                        unsafe_allow_html=True
                    )

                # Add copy button
                if st.button("Copy Summary to Clipboard"):
                    st.code(summary)
                    st.info("You can now copy the text above")

                st.markdown('</div>', unsafe_allow_html=True)

# Tab 2: Document Comparison Feature
# Tab 2: Document Comparison Feature
with tab2:
    st.header("Document Comparison")
    st.write("Upload two documents to compare their summaries side by side.")

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Document 1")
        doc1 = st.file_uploader("Upload first document", key="file1", type=["pdf", "txt", "docx"])

    with col2:
        st.subheader("Document 2")
        doc2 = st.file_uploader("Upload second document", key="file2", type=["pdf", "txt", "docx"])

    comparison_type = st.radio("Comparison Detail Level:", ["Brief", "Detailed"])

    if st.button("Compare Documents", key="compare_docs"):
        if not api_loaded:
            st.error("API key not loaded. Please check your configuration in KEYS/secrets.toml")
        elif doc1 is None or doc2 is None:
            st.warning("Please upload both documents to compare.")
        else:
            with st.spinner("Processing documents..."):
                try:
                    # Process Document 1
                    text1 = ""
                    if doc1.type == "application/pdf":
                        with fitz.open(stream=doc1.read(), filetype="pdf") as doc:
                            for page in doc:
                                text1 += page.get_text()
                    elif doc1.type == "text/plain":
                        text1 = doc1.getvalue().decode("utf-8")
                    elif doc1.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        doc = docx.Document(io.BytesIO(doc1.getvalue()))
                        text1 = "\n".join([paragraph.text for paragraph in doc.paragraphs])

                    # Process Document 2
                    text2 = ""
                    if doc2.type == "application/pdf":
                        with fitz.open(stream=doc2.read(), filetype="pdf") as doc:
                            for page in doc:
                                text2 += page.get_text()
                    elif doc2.type == "text/plain":
                        text2 = doc2.getvalue().decode("utf-8")
                    elif doc2.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        doc = docx.Document(io.BytesIO(doc2.getvalue()))
                        text2 = "\n".join([paragraph.text for paragraph in doc.paragraphs])


                    # Manual text splitting instead of NLTK
                    def simple_split_text(text, max_length=4000):
                        """Simple function to split text without NLTK dependencies"""
                        # If text is short enough, return as is
                        if len(text) <= max_length:
                            return [text]

                        # Otherwise, find sentence boundaries to split on
                        chunks = []
                        current_chunk = ""

                        # Basic sentence splitting by punctuation
                        sentences = re.split(r'(?<=[.!?])\s+', text)

                        for sentence in sentences:
                            if len(current_chunk) + len(sentence) + 1 <= max_length:
                                if current_chunk:
                                    current_chunk += " " + sentence
                                else:
                                    current_chunk = sentence
                            else:
                                chunks.append(current_chunk)
                                current_chunk = sentence

                        if current_chunk:
                            chunks.append(current_chunk)

                        return chunks


                    # Check documents aren't too short
                    if len(text1) < 10 or len(text2) < 10:
                        st.error("One or both documents are too short or empty. Please check your files.")
                        st.stop()


                    # Generate summaries using simple splitting to avoid NLTK
                    # Create a modified get_summary function for comparison
                    def get_comparison_summary(text, summary_type, model="gpt-4"):
                        """Modified summary function that doesn't rely on NLTK"""
                        if not text or len(text) < 10:
                            return "Error: Text is too short to summarize", 0

                        start_time = time.time()

                        prompt_template = "Write a concise summary of the following text."

                        # Check if text needs to be split
                        chunks = simple_split_text(text)

                        if len(chunks) == 1:
                            # Single chunk processing
                            prompt = f"{prompt_template}\n\n{text}"

                            try:
                                response = client.chat.completions.create(
                                    model=model,
                                    messages=[
                                        {"role": "system",
                                         "content": "You are a professional summarizer that creates clear, concise, and accurate summaries."},
                                        {"role": "user", "content": prompt}
                                    ],
                                    temperature=0.3,
                                    max_tokens=600
                                )

                                summary = response.choices[0].message.content.strip()

                            except Exception as e:
                                return f"Error generating summary: {str(e)}", 0

                        else:
                            # Multi-chunk processing for large documents
                            chunk_summaries = []

                            for chunk in chunks:
                                chunk_prompt = f"Summarize the following section of a document:\n\n{chunk}"

                                try:
                                    response = client.chat.completions.create(
                                        model=model,
                                        messages=[
                                            {"role": "system",
                                             "content": "You are a professional summarizer that creates clear, concise, and accurate summaries."},
                                            {"role": "user", "content": chunk_prompt}
                                        ],
                                        temperature=0.3,
                                        max_tokens=300
                                    )

                                    chunk_summary = response.choices[0].message.content.strip()
                                    chunk_summaries.append(chunk_summary)

                                except Exception as e:
                                    chunk_summaries.append(f"[Error summarizing chunk: {str(e)}]")

                            # Combine and summarize the chunk summaries
                            combined_text = "\n\n".join(chunk_summaries)
                            final_prompt = f"{prompt_template}\n\nBelow are summaries of different sections of a document. Create a unified summary:\n\n{combined_text}"

                            try:
                                response = client.chat.completions.create(
                                    model=model,
                                    messages=[
                                        {"role": "system",
                                         "content": "You are a professional summarizer that creates clear, concise, and accurate summaries."},
                                        {"role": "user", "content": final_prompt}
                                    ],
                                    temperature=0.3,
                                    max_tokens=600
                                )

                                summary = response.choices[0].message.content.strip()

                            except Exception as e:
                                return f"Error in final summarization: {str(e)}", 0

                        processing_time = time.time() - start_time

                        return summary, processing_time


                    # Generate summaries
                    progress_bar = st.progress(0)
                    st.info("Generating summary for document 1...")
                    summary1, _ = get_comparison_summary(text1, "Medium", model=model)

                    progress_bar.progress(0.5)
                    st.info("Generating summary for document 2...")
                    summary2, _ = get_comparison_summary(text2, "Medium", model=model)

                    progress_bar.progress(0.8)

                    # Check for errors in summaries
                    if "Error" in summary1 or "Error" in summary2:
                        if "Error" in summary1:
                            st.error(f"Error summarizing document 1: {summary1}")
                        if "Error" in summary2:
                            st.error(f"Error summarizing document 2: {summary2}")
                        st.stop()

                    # Generate comparison
                    st.info("Comparing documents...")
                    comparison_prompt = f"""Compare the following two summaries:

                    Summary 1:
                    {summary1}

                    Summary 2:
                    {summary2}

                    Provide a {'detailed' if comparison_type == 'Detailed' else 'brief'} analysis of the key similarities and differences between these documents.
                    """

                    response = client.chat.completions.create(
                        model=model,
                        messages=[
                            {"role": "system", "content": "You are a professional document analyst."},
                            {"role": "user", "content": comparison_prompt}
                        ],
                        temperature=0.5,
                        max_tokens=800
                    )

                    comparison = response.choices[0].message.content.strip()
                    progress_bar.progress(1.0)

                    # Display results
                    st.subheader("Comparison Results")

                    tabs = st.tabs(["Comparison", "Document 1 Summary", "Document 2 Summary"])

                    with tabs[0]:
                        st.write(comparison)

                    with tabs[1]:
                        st.subheader(f"Summary of {doc1.name}")
                        st.write(summary1)

                    with tabs[2]:
                        st.subheader(f"Summary of {doc2.name}")
                        st.write(summary2)

                except Exception as e:
                    st.error(f"Error comparing documents: {str(e)}")
                    # Add more detailed error information for debugging
                    import traceback

                    st.error(f"Detailed error: {traceback.format_exc()}")

# Tab 3: Help & Examples
with tab3:
    st.header("Help & Examples")

    with st.expander("How to Use This Tool"):
        st.write("""
        ### Getting Started
        1. Select your desired summary type (Short, Medium, or Detailed)
        2. Choose to either paste text or upload a document
        3. Click "Generate Summary" to process your text

        ### Advanced Features
        - **Custom Prompts**: Enable custom prompts to tailor your summary for specific purposes
        - **Document Comparison**: Compare two documents to identify similarities and differences
        - **History Tracking**: Access your previous summaries from the sidebar

        ### Supported File Types
        - PDF (.pdf)
        - Plain Text (.txt)
        - Word Documents (.docx)
        - CSV Data (.csv)
        """)

    with st.expander("Example Texts"):
        example_options = {
            "Data Science Report": "In this study, we applied machine learning techniques to predict housing prices in the San Francisco Bay Area. Using a dataset containing over 50,000 records of housing sales, we employed linear regression, decision trees, and random forests to build predictive models. The random forest model outperformed others with an R-squared value of 0.92. We also used feature engineering to create new variables, including the proximity to public transport, which significantly improved model accuracy. Model evaluation was performed using cross-validation with 10 folds, ensuring robustness in our predictions. We conclude that machine learning models can significantly enhance real estate pricing predictions and provide valuable insights for both buyers and sellers in competitive markets.",

            "Legal Brief": "This memorandum outlines the legal considerations in the case of Smith v. Johnson, where the plaintiff claims breach of contract related to the non-delivery of goods as specified in a signed agreement. The contract stipulated delivery within 30 days, but the defendant failed to meet this deadline, resulting in significant financial loss for the plaintiff. Under contract law, failure to perform obligations as per the agreed terms constitutes a breach. The defendant argues that delays were caused by unforeseen circumstances, which could be considered a valid defense under force majeure clauses. We recommend a thorough examination of the contract’s force majeure provisions, as well as gathering evidence to establish the extent of the plaintiff’s damages. The case could potentially be settled through mediation, although litigation remains a possibility if no agreement is reached.",

            "Finance Report": "In Q1 2025, the financial performance of the company exceeded expectations, with a 12% increase in earnings per share (EPS) compared to Q1 2024. This growth was driven by a combination of cost-cutting measures, increased demand in the consumer electronics sector, and strategic investments in emerging markets. However, the company’s debt-to-equity ratio has increased slightly to 1.4, which could raise concerns among investors regarding leverage. The finance team is recommending a focus on reducing debt through targeted capital allocation strategies. Additionally, the company plans to issue a dividend payout of $0.50 per share, up from $0.45 last quarter, signaling confidence in the company's continued financial stability. We also note a 6% decline in the stock price, which analysts suggest may be attributed to external macroeconomic factors such as interest rate hikes and inflationary pressures."
        }

        selected_example = st.selectbox("Select an example text", list(example_options.keys()))

        st.write("### Example Text Preview:")
        st.write(example_options[selected_example][:500] + "...")

        if st.button("Use This Example"):
            st.session_state.current_text = example_options[selected_example]
            st.rerun()

    with st.expander("Tips for Better Summaries"):
        st.write("""
        ### Tips for Better Results

        1. **Provide Clean Text**: Remove headers, footers, and irrelevant content for better summaries

        2. **Choose the Right Summary Type**:
           - Use **Short** for quick overviews
           - Use **Medium** for balanced summaries
           - Use **Detailed** for comprehensive analysis

        3. **Try Different Models**:
           - GPT-3.5 is faster but less detailed
           - GPT-4 provides higher quality summaries

        4. **Use Custom Prompts** to focus summaries on specific aspects of your text
        """)

    with st.expander("FAQ"):
        st.write("""
        ### Frequently Asked Questions

        **Q: How long can my document be?**  
        A: The tool works best with documents under 30,000 words. Longer documents will be automatically chunked for processing.

        **Q: Are my documents stored?**  
        A: Documents are processed in memory and not permanently stored. Summaries are saved only in your browser session.

        **Q: What languages are supported?**  
        A: The tool translates foreign languages to English Summary.

        **Q: Can I download my summaries?**  
        A: Yes, download options are provided for both the original text and the generated summaries.

        **Q: How accurate are the summaries?**  
        A: The AI models are trained to extract key information, but you should always review the generated summaries for critical applications.
        """)

# Footer with credits
st.markdown("""
<div class="footer">
    <p>AI-Powered Text Summarizer | GEN AI Final Project | © 2025</p>
</div>
""", unsafe_allow_html=True)


# Additional functions for utility purposes

def extract_key_terms(text, num_terms=10):
    """Extract key terms from text using basic frequency analysis"""
    import re
    from collections import Counter

    # Remove common punctuation and convert to lowercase
    text = re.sub(r'[^\w\s]', '', text.lower())

    # Remove common stopwords
    stopwords = {'a', 'an', 'the', 'and', 'or', 'but', 'if', 'because', 'as', 'what',
                 'with', 'by', 'for', 'is', 'in', 'to', 'that', 'of', 'it', 'this',
                 'be', 'are', 'was', 'were', 'been', 'being', 'have', 'has', 'had',
                 'do', 'does', 'did', 'can', 'could', 'will', 'would', 'should', 'on',
                 'at', 'from', 'about', 'them', 'these', 'those', 'they', 'we', 'he',
                 'she', 'his', 'her', 'their', 'our', 'your', 'my', 'mine', 'yours',
                 'theirs', 'ours', 'i', 'you', 'not', 'no', 'yes', 'then', 'than'}

    words = [word for word in text.split() if word not in stopwords and len(word) > 3]

    # Count word frequency
    word_counts = Counter(words)

    # Get the most common words
    return word_counts.most_common(num_terms)


def create_word_cloud(terms):
    """Create a simple word cloud visualization using plotly"""
    words = [term[0] for term in terms]
    counts = [term[1] for term in terms]

    # Create a bubble chart for word cloud
    fig = go.Figure(data=[go.Scatter(
        x=[i for i in range(len(words))],
        y=[1 for _ in range(len(words))],
        mode='text+markers',
        text=words,
        marker=dict(
            size=[c * 20 for c in counts],
            color=[c for c in counts],
            colorscale='Viridis',
            showscale=True,
            colorbar=dict(title="Frequency")
        ),
        textfont=dict(
            size=[c * 5 for c in counts],
            color='white'
        )
    )])

    fig.update_layout(
        title="Key Terms Word Cloud",
        xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        yaxis=dict(showgrid=False, zeroline=False, showticklabels=False)
    )

    return fig


def get_differences(text1, text2):
    """Generate a comparison of two texts using the API"""
    prompt = f"""Compare the following two texts and identify key differences:

    Text 1:
    {text1[:1000]}

    Text 2:
    {text2[:1000]}

    List the main differences in terms of content, focus, and key points:
    """

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",  # Using the faster model for comparisons
            messages=[
                {"role": "system", "content": "You are a text analysis assistant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=400
        )

        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error generating comparison: {str(e)}"